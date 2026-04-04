"""
Vector-Based RAG Pipeline using Chroma
========================================
Stack:
  LLM        : Google Gemini (GOOGLE_API_KEY)
  Embeddings : BAAI/bge-small-en-v1.5 (HuggingFace)
  Vector DB  : Chroma (persistent local storage)
  Framework  : LangChain

Supported file types: PDF, DOCX, XLSX, TXT
"""

import os
import re
import hashlib
from fastapi import File
import numpy as np
from pathlib import Path
from typing import List, Dict, Optional, Tuple

# ── Load System Prompt from external file ──────────────────────────────────────
PROMPT_FILE = Path(__file__).parent / "system_prompt.txt"
if PROMPT_FILE.exists():
    SYSTEM_PROMPT = PROMPT_FILE.read_text(encoding="utf-8")
else:
    print(f"[Warning] system_prompt.txt not found at {PROMPT_FILE}")
    SYSTEM_PROMPT = "You are a precise document question-answering assistant."

# ── LangChain imports ──────────────────────────────────────────────────────────
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_chroma import Chroma
from langchain_community.embeddings import HuggingFaceEmbeddings


# ══════════════════════════════════════════════════════════════════════════════
# 1. DOCUMENT LOADERS
# ══════════════════════════════════════════════════════════════════════════════

def load_pdf(path: str) -> List[Document]:
    """PyMuPDF — preserves page numbers."""
    import fitz
    docs = []
    with fitz.open(path) as pdf:
        for i, page in enumerate(pdf, 1):
            text = page.get_text("text").strip()
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={"source": Path(path).name, "page": i, "type": "pdf"}
                ))
    return docs


def load_docx(path: str) -> List[Document]:
    """python-docx — groups paragraphs into virtual pages (~40 per page)."""
    from docx import Document as DocxDoc
    raw = DocxDoc(path)
    paras = [p.text.strip() for p in raw.paragraphs if p.text.strip()]
    docs, page_size = [], 40
    for i in range(0, len(paras), page_size):
        block = "\n".join(paras[i:i + page_size])
        docs.append(Document(
            page_content=block,
            metadata={"source": Path(path).name, "page": (i // page_size) + 1, "type": "docx"}
        ))
    return docs


def load_xlsx(path: str) -> List[Document]:
    """openpyxl — one Document per sheet, rows rendered as tab-separated text."""
    import openpyxl
    wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    docs = []
    for sheet_idx, ws in enumerate(wb.worksheets, 1):
        rows, headers = [], None
        for row in ws.iter_rows(values_only=True):
            if all(v is None for v in row):
                continue
            vals = [str(v) if v is not None else "" for v in row]
            if headers is None and any(vals):
                headers = vals
            else:
                rows.append(vals)
        if headers:
            lines = ["\t".join(headers)] + ["\t".join(r) for r in rows[:300]]
            docs.append(Document(
                page_content=f"Sheet: {ws.title}\n" + "\n".join(lines),
                metadata={"source": Path(path).name, "page": sheet_idx,
                          "sheet": ws.title, "type": "xlsx"}
            ))
    wb.close()
    return docs


def load_txt(path: str) -> List[Document]:
    """Plain text loader."""
    return [Document(
        page_content=Path(path).read_text(errors="ignore"),
        metadata={"source": Path(path).name, "page": 1, "type": "txt"}
    )]


LOADERS: Dict = {
    ".pdf": load_pdf, ".docx": load_docx,
    ".xlsx": load_xlsx, ".xls": load_xlsx, ".txt": load_txt,
}


def load_documents(file_paths: List[str]) -> List[Document]:
    """Load all files into a flat list of LangChain Documents."""
    all_docs = []
    for fp in file_paths:
        p = Path(fp)
        loader = LOADERS.get(p.suffix.lower())
        if not loader:
            print(f"[Loader] ⚠ Unsupported: {p.name}")
            continue
        try:
            docs = loader(str(p))
            print(f"[Loader] ✓ {p.name} → {len(docs)} page(s)")
            all_docs.extend(docs)
        except Exception as e:
            print(f"[Loader] ✗ {p.name}: {e}")
    return all_docs


# ══════════════════════════════════════════════════════════════════════════════
# 2. EMBEDDING ENGINE
# ══════════════════════════════════════════════════════════════════════════════

class QwenEmbeddingEngine:
    """
    HuggingFace Embedding Engine using sentence-transformers.
    Uses BAAI/bge-small-en-v1.5 for faster local embeddings (no API needed).
    Falls back to sentence-transformers if specified model unavailable.
    """

    MODEL = "BAAI/bge-small-en-v1.5"

    def __init__(self, api_key: Optional[str] = None, batch_size: int = 32):
        self.batch_size = batch_size
        print(f"[Embeddings] Loading {self.MODEL}...")
        try:
            self._embedder = HuggingFaceEmbeddings(
                model_name=self.MODEL,
                encode_kwargs={"normalize_embeddings": True},
                show_progress=True,
            )
            print(f"[Embeddings] ✓ Model loaded: {self.MODEL}")
        except Exception as e:
            print(f"[Embeddings] ⚠ Warning: Could not load {self.MODEL}: {e}")
            print(f"[Embeddings] Trying fallback model...")
            self._embedder = HuggingFaceEmbeddings(
                model_name="sentence-transformers/all-MiniLM-L6-v2",
                encode_kwargs={"normalize_embeddings": True},
                show_progress=True,
            )
            print(f"[Embeddings] ✓ Using fallback: all-MiniLM-L6-v2")

    def embed_documents(self, texts: List[str]) -> np.ndarray:
        """Embed a list of passage texts → (N, D) float32 array, L2-normalised."""
        try:
            vecs = self._embedder.embed_documents(texts)
            arr = np.array(vecs, dtype=np.float32)
            arr /= np.linalg.norm(arr, axis=1, keepdims=True) + 1e-10
            return arr
        except Exception as e:
            print(f"[Embeddings] Error embedding documents: {e}")
            raise

    def embed_query(self, text: str) -> np.ndarray:
        """Embed a single query → (D,) float32 array, L2-normalised."""
        try:
            vec = np.array(self._embedder.embed_query(text), dtype=np.float32)
            vec /= np.linalg.norm(vec) + 1e-10
            return vec
        except Exception as e:
            print(f"[Embeddings] Error embedding query: {e}")
            raise


# ══════════════════════════════════════════════════════════════════════════════
# 3. VECTOR-BASED RAG PIPELINE
# ══════════════════════════════════════════════════════════════════════════════

QA_PROMPT = ChatPromptTemplate.from_messages([
    ("system", SYSTEM_PROMPT),
    ("human", """Sources:
{context}

Question: {question}

Answer with inline [SOURCE-N] citations:"""),
])


class VectorRAG:
    """
    Full LangChain RAG pipeline backed by Chroma vector database.

    Usage:
        rag = VectorRAG(persist_dir="./chroma_db")
        rag.ingest(["policy.pdf", "data.xlsx"])
        result = rag.query("What is the ambient pallet storage rate?")
    """

    def __init__(
        self,
        gemini_model: str = "gemini-2.5-flash",
        chunk_size: int = 600,
        chunk_overlap: int = 80,
        top_k: int = 8,
        persist_dir: str = "./chroma_db",
        google_api_key: Optional[str] = None,
        hf_api_key: Optional[str] = None,
    ):
        # ── LLM ────────────────────────────────────────────────────────────────
        g_key = google_api_key or os.environ.get("GOOGLE_API_KEY")
        if not g_key:
            raise ValueError("GOOGLE_API_KEY not set.")
        self.llm = ChatGoogleGenerativeAI(
            model=gemini_model,
            google_api_key=g_key,
            temperature=0.0,
        )
        print(f"[LLM] Using {gemini_model}")

        # ── Embeddings ─────────────────────────────────────────────────────────
        self.embedder = QwenEmbeddingEngine(api_key=hf_api_key)

        # ── Splitter ───────────────────────────────────────────────────────────
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

        # ── Vector DB (Chroma) ─────────────────────────────────────────────────
        self.persist_dir = persist_dir
        self.vectorstore = None
        self.top_k = top_k
        self._init_vectorstore()

    def _init_vectorstore(self):
        """Initialize or load existing Chroma vectorstore."""
        Path(self.persist_dir).mkdir(parents=True, exist_ok=True)
        self.vectorstore = Chroma(
            embedding_function=self.embedder._embedder,
            persist_directory=self.persist_dir,
            collection_name="meridian_docs"
        )
        print(f"[VectorDB] Chroma initialized at {self.persist_dir}")

    def ingest(self, files=File()) -> Dict:
        """
        Load → split → embed → index all documents into Chroma.
        Returns ingestion stats.
        """
        stats = {"files": 0, "pages": 0, "chunks": 0, "errors": []}

        raw_docs = load_documents(file_paths)
        stats["pages"] = len(raw_docs)

        if not raw_docs:
            print("[Ingest] No documents loaded.")
            return stats

        # Split each page-doc into chunks, preserving metadata
        all_lc_chunks: List[Document] = []
        for doc in raw_docs:
            splits = self.splitter.split_documents([doc])
            for idx, split in enumerate(splits):
                split.metadata["chunk_index"] = idx
            all_lc_chunks.extend(splits)

        print(f"[Ingest] {len(raw_docs)} pages → {len(all_lc_chunks)} chunks")

        # Add documents to Chroma
        print(f"[Ingest] Adding {len(all_lc_chunks)} chunks to Chroma...")
        
        ids = []
        for idx, chunk in enumerate(all_lc_chunks):
            uid_raw = f"{chunk.metadata.get('source')}|{chunk.metadata.get('page')}|{idx}|{chunk.page_content[:64]}"
            chunk_id = hashlib.sha256(uid_raw.encode()).hexdigest()[:16]
            ids.append(chunk_id)
        
        try:
            self.vectorstore.add_documents(all_lc_chunks, ids=ids)
            print(f"[Ingest] ✓ Documents added to Chroma")
        except Exception as e:
            print(f"[Ingest] ✗ Error adding documents: {e}")
            raise
        
        # Persist to disk
        try:
            self.vectorstore.persist()
            print(f"[Ingest] ✓ Persisted to disk")
        except Exception as e:
            print(f"[Ingest] ⚠ Warning: Persist failed: {e}")
        
        # Reload vectorstore
        try:
            self._init_vectorstore()
            print(f"[Ingest] ✓ Reloaded vectorstore")
        except Exception as e:
            print(f"[Ingest] ⚠ Warning: Reload failed: {e}")

        stats["files"] = len(set([c.metadata.get("source") for c in all_lc_chunks]))
        stats["chunks"] = len(all_lc_chunks)
        print(f"[Ingest] ✓ Indexed {stats['chunks']} chunks from {stats['files']} files")

        return stats

    def retrieve(self, question: str) -> List[Tuple[Document, float]]:
        """Retrieve top-k chunks with similarity scores from Chroma."""
        results = self.vectorstore.similarity_search_with_score(question, k=self.top_k)
        return results

    def query(self, question: str) -> Dict:
        """
        Run full RAG: retrieve → build context → Gemini → parse citations.
        """
        results = self.retrieve(question)
        if not results:
            return {"question": question, "answer": "No documents indexed.", "sources": [], "retrieved": []}

        context_blocks = []
        source_registry: Dict[str, Dict] = {}
        
        for rank, (doc, score) in enumerate(results, 1):
            label = f"[SOURCE-{rank}]"
            header = f"{label} — {doc.metadata.get('source')}, Page {doc.metadata.get('page')} (score: {score:.4f})"
            context_blocks.append(f"{header}\n{doc.page_content}")
            source_registry[label] = {
                "label": label,
                "doc_id": doc.metadata.get("source"),
                "page_num": doc.metadata.get("page"),
                "chunk_index": doc.metadata.get("chunk_index", 0),
                "score": round(float(score), 4),
                "preview": doc.page_content[:200] + ("…" if len(doc.page_content) > 200 else ""),
            }

        context = "\n\n---\n\n".join(context_blocks)

        # Call Gemini via LangChain
        formatted = QA_PROMPT.format_messages(context=context, question=question)
        answer = (self.llm | StrOutputParser()).invoke(formatted)

        # Extract cited source labels
        cited_labels = sorted(set(re.findall(r'\[SOURCE-\d+\]', answer)))
        cited_sources = [source_registry[l] for l in cited_labels if l in source_registry]

        return {
            "question": question,
            "answer": answer,
            "sources": cited_sources,
            "retrieved": list(source_registry.values()),
        }

    def get_stats(self) -> Dict:
        """Get statistics about the vectorstore."""
        try:
            collection = self.vectorstore._collection
            return {"total_documents": collection.count()}
        except:
            return {"total_documents": "Unable to retrieve"}